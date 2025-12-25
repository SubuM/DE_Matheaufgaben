import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document
import tempfile
import io

# -----------------------
# Utility: Extract text
# -----------------------

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# -----------------------
# Streamlit App UI
# -----------------------

st.set_page_config(page_title="Interactive Math Question Generator", layout="wide")
st.title("📘 Interactive Grade 5 Math Question Generator")

st.write("""
Upload the **official syllabus** (PDF or DOCX).  
Then select **difficulty level**, **number of sets** (1–10), and each set will contain **50 random questions**.  
The app will generate a **DOCX file** with all sets, properly separated and numbered.
""")

uploaded_file = st.file_uploader("Upload syllabus PDF or DOCX", type=["pdf", "docx"])

# -----------------------
# Load API key from secrets
# -----------------------

if "groq" not in st.secrets or "api_key" not in st.secrets["groq"]:
    st.error("❌ GROQ API key missing in .streamlit/secrets.toml under [groq] section.")
    st.stop()

API_KEY = st.secrets["groq"]["api_key"]
client = Groq(api_key=API_KEY)

syllabus_text = ""

# -----------------------
# File Processing
# -----------------------

if uploaded_file:
    filetype = uploaded_file.name.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    if filetype == "pdf":
        st.info("Extracting text from PDF…")
        try:
            syllabus_text = extract_text_from_pdf(tmp_path)
        except Exception as e:
            st.error(f"Failed to extract PDF: {e}")

    elif filetype == "docx":
        st.info("Extracting text from DOCX…")
        try:
            syllabus_text = extract_text_from_docx(tmp_path)
        except Exception as e:
            st.error(f"Failed to extract DOCX: {e}")

    else:
        st.error("Unsupported file type.")

# -----------------------
# Question settings
# -----------------------

if syllabus_text:
    st.subheader("⚙️ Question Settings")

    num_sets = st.number_input(
        "Number of sets to generate (max 10)",
        min_value=1,
        max_value=10,
        value=1,
        step=1
    )

    questions_per_set = 50  # fixed per set

    difficulty = st.select_slider(
        "Difficulty level:",
        options=["Easy", "Medium", "Hard", "Very Hard"],
        value="Hard"
    )

    difficulty_explanations = {
        "Easy": "simple arithmetic and basic understanding",
        "Medium": "standard grade 5 problems with moderate reasoning",
        "Hard": "multi-step reasoning, word problems, challenging logic",
        "Very Hard": "highly challenging, advanced reasoning for grade 5"
    }

    st.write(f"**Difficulty selected:** {difficulty} — {difficulty_explanations[difficulty]}")
    st.write(f"**Each set will contain:** {questions_per_set} questions")

    if st.button("Generate Sets (DOCX)"):
        with st.spinner(f"Generating {num_sets} sets at {difficulty} difficulty…"):
            
            doc = Document()
            doc.add_heading(f"Interactive Math Question Sets", 0)
            doc.add_paragraph("Generated from Sachsen-Anhalt Gymnasium Grade 5 syllabus.\n")

            question_number_global = 1

            for set_idx in range(1, num_sets+1):
                doc.add_heading(f"Set {set_idx}", level=1)

                prompt = f"""
                You are an expert mathematics teacher.

                Based on the following **grade 5 Gymnasium syllabus for Sachsen-Anhalt**, 
                generate **exactly {questions_per_set} random math questions** for **Set {set_idx}**.

                Difficulty level: **{difficulty}**  
                Meaning: {difficulty_explanations[difficulty]}

                Requirements:
                - Use a variety of question types: word problems, multi-step reasoning,
                  geometry, number theory, fractions, arithmetic, logic puzzles,
                  real-world applications.
                - Cover all relevant topics from the syllabus.
                - Number questions sequentially starting from {question_number_global}.
                - Do NOT include solutions.
                - Randomize question order within this set.

                Syllabus content:
                {syllabus_text}
                """

                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}]
                )

                questions_text = response.choices[0].message.content

                # Add questions to DOCX
                for line in questions_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

                question_number_global += questions_per_set

            # Save DOCX to BytesIO
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.success(f"DOCX file ready with {num_sets} sets! 🎉")

            st.download_button(
                label="📥 Download All Sets (DOCX)",
                data=doc_io,
                file_name=f"math_question_sets_{difficulty}_{num_sets}sets.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
