import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document
import tempfile
import io

# -----------------------
# Hilfsfunktionen
# -----------------------

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# -----------------------
# UI
# -----------------------

st.set_page_config(
    page_title="Mathematik-Aufgabengenerator Klasse 5/6 (Gymnasium)",
    layout="wide"
)

st.title("📘 Anspruchsvoller Mathematik-Aufgabengenerator (Gymnasium)")

st.write("""
Dieser Generator erstellt **anspruchsvolle Mathematikaufgaben**  
auf Basis des offiziellen **Lehrplans Klasse 5 (Gymnasium Sachsen-Anhalt)**.

🔹 **Schwierigkeitsniveau:** mindestens **Klasse 6**, teilweise darüber  
🔹 **Hohe Zufälligkeit** innerhalb aller Lehrplanbereiche  
🔹 **KMK-konformer Schulbuchstil**  
🔹 **Keine Lösungen**
""")

uploaded_file = st.file_uploader(
    "📤 Lehrplan hochladen (PDF oder DOCX)",
    type=["pdf", "docx"]
)

# -----------------------
# GROQ
# -----------------------

if "groq" not in st.secrets or "api_key" not in st.secrets["groq"]:
    st.error("❌ GROQ API-Schlüssel fehlt.")
    st.stop()

client = Groq(api_key=st.secrets["groq"]["api_key"])
syllabus_text = ""

# -----------------------
# Datei lesen
# -----------------------

if uploaded_file:
    filetype = uploaded_file.name.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    if filetype == "pdf":
        syllabus_text = extract_text_from_pdf(tmp_path)
    elif filetype == "docx":
        syllabus_text = extract_text_from_docx(tmp_path)

# -----------------------
# Einstellungen
# -----------------------

if syllabus_text:
    st.subheader("⚙️ Einstellungen")

    num_sets = st.number_input(
        "Anzahl der Aufgabensätze",
        min_value=1,
        max_value=10,
        value=1
    )

    questions_per_set = 50

    difficulty = st.select_slider(
        "Formaler Schwierigkeitsgrad:",
        options=["Schwer", "Sehr schwer"],
        value="Sehr schwer"
    )

    difficulty_explanations = {
        "Schwer": "anspruchsvolle Aufgaben mit mehreren Denk- und Rechenschritten",
        "Sehr schwer": "hohe kognitive Anforderungen, Vergleich, Begründung, Transfer"
    }

    st.write("🔍 **Hinweis:** Inhalt Klasse 5, Denkanspruch mindestens Klasse 6")

    # -----------------------
    # Generieren
    # -----------------------

    if st.button("🔥 Anspruchsvolle Aufgabensätze generieren"):
        with st.spinner("Aufgaben werden erstellt …"):

            doc = Document()
            doc.add_heading("Mathematik – Anspruchsvolle Aufgabensätze", 0)
            doc.add_paragraph(
                "Lehrplan Klasse 5 (Gymnasium Sachsen-Anhalt)\n"
                "Anforderungsniveau: Klasse 6 / erhöhte Kompetenzstufe\n"
            )

            question_number_global = 1

            for set_idx in range(1, num_sets + 1):
                doc.add_heading(f"Aufgabensatz {set_idx}", level=1)

                prompt = f"""
Du bist ein sehr erfahrener deutscher Mathematiklehrer am Gymnasium
mit Schwerpunkt auf leistungsstarken Lerngruppen.

Erstelle **genau {questions_per_set} unterschiedliche Mathematikaufgaben**
auf Grundlage des folgenden **Lehrplans Klasse 5 (Gymnasium Sachsen-Anhalt)**.

🎯 ZIEL:
- Inhaltlich Klasse 5
- **Denk- und Anspruchsniveau mindestens Klasse 6**

🔀 ZUFÄLLIGKEIT (sehr wichtig):
- Verteile die Aufgaben **zufällig und ausgewogen** auf alle Lehrplanbereiche
- Verwende **unterschiedliche Zahlenräume, Kontexte, Darstellungen**
- Vermeide erkennbare Muster oder Wiederholungen
- Jede Aufgabe soll sich klar von den anderen unterscheiden

📘 DIDAKTISCHE VORGABEN (KMK-Stil):
- Schulbuchnahe, präzise Formulierungen
- Klare Arbeitsaufträge
- Häufig mehrschrittige Lösungswege erforderlich
- Vergleichs-, Begründungs- und Transferaufgaben einbauen

📌 AUFGABENTYPEN (mischen):
- Anspruchsvolle Textaufgaben
- Mehrschrittige Rechnungen
- Geometrische Denkaufgaben
- Brüche & natürliche Zahlen kombiniert
- Sachprobleme mit Auswahl relevanter Informationen
- Logische Schlussfolgerungen und Vergleiche

🚫 STRIKT:
- **Keine Lösungen**
- **Keine Hinweise**
- **Keine Zwischenschritte**

🔢 Nummerierung:
- Fortlaufend ab {question_number_global}

📚 Lehrplan:
{syllabus_text}
"""

                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}]
                )

                questions_text = response.choices[0].message.content

                for line in questions_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

                question_number_global += questions_per_set

            # -----------------------
            # Speichern
            # -----------------------

            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            st.success("✅ Anspruchsvolle Aufgabensätze erstellt!")

            st.download_button(
                label="📥 DOCX herunterladen",
                data=output,
                file_name=f"mathematik_anspruchsvoll_klasse5_6_{num_sets}_sets.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
