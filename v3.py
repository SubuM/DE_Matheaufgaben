import streamlit as st
import random
from docx import Document
from io import BytesIO
import zipfile
import os

# --- Generator Functions (Tough Problems - Omitted for brevity, assume they are still here) ---

# --- NOTE: The GENERATORS list and all tough generation functions would be defined here ---
#           For this response, we'll define a simple placeholder to make the code runnable.
GENERATORS = [lambda: "Berechne: 10 + 5 x 2", lambda: "Runde 987654 auf Tausender"] 
# Ensure the full tough generation functions are included in your actual file.


def create_single_problem_set(num_problems=50):
    problems = []
    min_per_category = 7 
    required_problems = []
    
    # ... (Rest of the problem generation logic)
    
    # Simple placeholder logic for demonstrating the change:
    for i in range(1, num_problems + 1):
        problem = random.choice(GENERATORS)()
        problems.append(f"{i}. {problem}")
        
    return problems


def create_word_document(problems, set_number):
    """Erstellt ein Word-Dokument (im Speicher) und gibt es als Bytes zurück."""
    document = Document()
    document.add_heading(f'Schwere Mathematikaufgaben Gymnasium Kl. 5 (Sachsen-Anhalt) - Set {set_number}', 0)
    document.add_paragraph("Dies sind Übungen mit erhöhtem Schwierigkeitsgrad.")
    for problem in problems:
        document.add_paragraph(problem)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- 3. Streamlit Login und Hauptprogramm ---

def login_form():
    """Zeigt das Login-Formular an und prüft die Anmeldedaten aus st.secrets."""
    
    # Lese die Anmeldedaten aus der secrets.toml Datei
    try:
        STATIC_USER = st.secrets["auth"]["username"]
        STATIC_PASS = st.secrets["auth"]["password"]
    except KeyError:
        st.error("Konfigurationsfehler: 'secrets.toml' oder die Sektion [auth] fehlt.")
        return

    st.sidebar.header("Login erforderlich")
    with st.sidebar.form("login_form"):
        username = st.text_input("Benutzername")
        password = st.text_input("Passwort", type="password")
        submitted = st.form_submit_button("Login")
        
        if submitted:
            if username == STATIC_USER and password == STATIC_PASS:
                st.session_state["authenticated"] = True
                st.success("Login erfolgreich!")
                # Verwenden Sie st.rerun() nur, wenn nötig, um den Zustand zu aktualisieren
                # Wenn Streamlit dies automatisch macht, kann es weggelassen werden
                # st.rerun() 
            else:
                st.error("Ungültiger Benutzername oder Passwort.")

def main_app():
    """Die Hauptanwendung, die nur nach erfolgreichem Login angezeigt wird."""
    st.title("💪 Matheaufgaben Generator (Hoher Schwierigkeitsgrad)")
    st.markdown("Generiert **10 völlig zufällige und schwere** Übungsblätter.")
    st.markdown("---")

    # User Inputs
    num_sets = st.number_input(
        "1. Anzahl der benötigten Aufgabensätze (Maximal 10 Sets empfohlen):", 
        min_value=1, 
        max_value=10, 
        value=5,
        help="Wie viele verschiedene Sätze von 50 schweren Aufgaben benötigen Sie?"
    )
    
    download_location = st.text_input(
        "2. Gewünschter lokaler Speicherpfad (zur Information):",
        value="C:/Users/IhrName/Downloads/Matheaufgaben_Schwer/",
        help="Der Browser kann Dateien nicht direkt dorthin speichern. Der Pfad dient nur als Hinweis."
    )
    
    st.markdown("---")

    if st.button(f"Starte Generierung von {num_sets} Sätzen und erstelle ZIP-Datei"):
        
        zip_buffer = BytesIO()
        status_placeholder = st.empty()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            
            with st.spinner(f"Generiere {num_sets} Aufgabensätze und komprimiere Dateien..."):
                
                for i in range(1, int(num_sets) + 1):
                    problems = create_single_problem_set(num_problems=50)
                    docx_bytes = create_word_document(problems, i)
                    filename = f"Matheaufgaben_Set_SCHWER_{i}.docx"
                    zf.writestr(filename, docx_bytes)
                    
                status_placeholder.success(f"✅ {num_sets} Aufgabensätze wurden erfolgreich generiert und zur ZIP-Datei hinzugefügt.")
                
        zip_buffer.seek(0)
        
        st.markdown("---")
        st.markdown(f"**ℹ️ Hinweis:** Die Dateien wären lokal unter dem Pfad: `{download_location}` gespeichert worden.")
        
        st.download_button(
            label="Alle Sätze als ZIP-Datei herunterladen",
            data=zip_buffer,
            file_name="Matheaufgaben_Klasse_5_SCHWER.zip",
            mime="application/zip",
            key='download_zip_button'
        )

# Hauptfunktion, die den Login-Zustand verwaltet
if __name__ == "__main__":
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"]:
        main_app()
    else:
        login_form()