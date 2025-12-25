import streamlit as st
from groq import Groq
import PyPDF2
from docx import Document
import tempfile
import io

# -----------------------
# Hilfsfunktionen: Text extrahieren
# -----------------------

def extract_text_from_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return text

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

# -----------------------
# Streamlit UI
# -----------------------

st.set_page_config(
    page_title="Interaktiver Mathematik-Aufgabengenerator Klasse 5",
    layout="wide"
)

st.title("📘 Interaktiver Mathematik-Aufgabengenerator – Klasse 5 (Gymnasium)")

st.write("""
Lade den **offiziellen Lehrplan** (PDF oder DOCX) hoch.  
Wähle anschließend den **Schwierigkeitsgrad** und die **Anzahl der Aufgabensätze** (1–10).

👉 Jeder Aufgabensatz enthält **genau 50 zufällige Aufgaben**.  
📄 Am Ende erhältst du **eine DOCX-Datei** mit allen Aufgabensätzen, sauber getrennt und nummeriert.
""")

uploaded_file = st.file_uploader(
    "📤 Lehrplan hochladen (PDF oder DOCX)",
    type=["pdf", "docx"]
)

# -----------------------
# GROQ API-Key laden
# -----------------------

if "groq" not in st.secrets or "api_key" not in st.secrets["groq"]:
    st.error("❌ GROQ API-Schlüssel fehlt in `.streamlit/secrets.toml` unter [groq].")
    st.stop()

API_KEY = st.secrets["groq"]["api_key"]
client = Groq(api_key=API_KEY)

syllabus_text = ""

# -----------------------
# Datei verarbeiten
# -----------------------

if uploaded_file:
    filetype = uploaded_file.name.split(".")[-1].lower()

    with tempfile.NamedTemporaryFile(delete=False) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    if filetype == "pdf":
        st.info("📄 Text wird aus dem PDF extrahiert …")
        syllabus_text = extract_text_from_pdf(tmp_path)

    elif filetype == "docx":
        st.info("📄 Text wird aus dem DOCX extrahiert …")
        syllabus_text = extract_text_from_docx(tmp_path)

    else:
        st.error("❌ Nicht unterstütztes Dateiformat.")

# -----------------------
# Aufgabeneinstellungen
# -----------------------

if syllabus_text:
    st.subheader("⚙️ Aufgabeneinstellungen")

    num_sets = st.number_input(
        "Anzahl der zu erstellenden Aufgabensätze (max. 10)",
        min_value=1,
        max_value=10,
        value=1,
        step=1
    )

    questions_per_set = 50  # fest

    difficulty = st.select_slider(
        "Schwierigkeitsgrad:",
        options=["Leicht", "Mittel", "Schwer", "Sehr schwer"],
        value="Schwer"
    )

    difficulty_explanations = {
        "Leicht": "einfache Rechenaufgaben und grundlegendes Verständnis",
        "Mittel": "typische Aufgaben der Klasse 5 mit einfachem Begründen",
        "Schwer": "mehrschrittige Denkaufgaben, anspruchsvolle Textaufgaben",
        "Sehr schwer": "sehr anspruchsvolle Aufgaben mit vertieftem logischem Denken"
    }

    st.write(f"**Gewählter Schwierigkeitsgrad:** {difficulty}")
    st.write(f"**Erläuterung:** {difficulty_explanations[difficulty]}")
    st.write(f"**Aufgaben pro Aufgabensatz:** {questions_per_set}")

    # -----------------------
    # Generierung
    # -----------------------

    if st.button("📘 Aufgabensätze generieren (DOCX)"):
        with st.spinner("✏️ Aufgabensätze werden erstellt …"):

            doc = Document()
            doc.add_heading("Mathematik – Aufgabensätze Klasse 5", 0)
            doc.add_paragraph(
                "Erstellt gemäß Lehrplan Gymnasium Sachsen-Anhalt "
                "und den Kompetenzformulierungen der KMK.\n"
            )

            question_number_global = 1

            for set_idx in range(1, num_sets + 1):
                doc.add_heading(f"Aufgabensatz {set_idx}", level=1)

                prompt = f"""
Du bist ein erfahrener deutscher Mathematiklehrer am Gymnasium.

Erstelle auf Grundlage des folgenden **Lehrplans für die Jahrgangsstufe 5
(Gymnasium Sachsen-Anhalt)** **genau {questions_per_set} Mathematikaufgaben**
für **Aufgabensatz {set_idx}**.

Die Aufgaben müssen:
- vollständig **auf Deutsch** formuliert sein
- sprachlich altersgerecht für Klasse 5 sein
- dem **KMK-typischen Schulbuchstil** entsprechen
- klar, sachlich und präzise formuliert sein

Schwierigkeitsgrad: **{difficulty}**  
Bedeutung: {difficulty_explanations[difficulty]}

Didaktische Vorgaben:
- Mischung aus: Rechenaufgaben, Textaufgaben, Geometrie,
  Brüche, natürliche Zahlen, Sachaufgaben,
  logisches Denken und einfache Begründungen
- Alle relevanten Inhalte des Lehrplans berücksichtigen
- Aufgaben **fortlaufend nummerieren**, beginnend mit {question_number_global}
- **Keine Lösungen, keine Hinweise, keine Zwischenschritte**
- Reihenfolge der Aufgaben zufällig wählen

Lehrplan:
{syllabus_text}
"""

                response = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=[{"role": "user", "content": prompt}]
                )

                questions_text = response.choices[0].message.content

                for line in questions_text.split("\n"):
                    if line.strip():
                        doc.add_paragraph(line.strip())

                question_number_global += questions_per_set

            # -----------------------
            # DOCX speichern
            # -----------------------

            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            st.success("✅ DOCX-Datei erfolgreich erstellt!")

            st.download_button(
                label="📥 Aufgabensätze herunterladen (DOCX)",
                data=doc_io,
                file_name=f"mathematik_klasse5_{difficulty}_{num_sets}_aufgabensaetze.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
