import streamlit as st
from openai import OpenAI
from docx import Document
from io import BytesIO
import zipfile
import time

# --- Configuration & Texts ---
TEXTS = {
    'de': {
        'title': "🧠 AI Mathe-Lehrer (Sachsen-Anhalt Geniessermodus)",
        'subtitle': "Nutzt künstliche Intelligenz, um den Lehrplan der 5. Klasse Gymnasium (Sachsen-Anhalt) dynamisch zu analysieren und schwere Aufgaben zu erstellen.",
        'api_label': "OpenAI API Key eingeben:",
        'api_help': "Der Key wird benötigt, um die KI-Modelle abzufragen. Er wird nicht gespeichert.",
        'sets_label': "Anzahl der Aufgabensätze:",
        'start_button': "🚀 Generiere Aufgaben mit AI",
        'status_gen': "🤖 Die KI analysiert den Lehrplan und generiert Set {}/{}...",
        'success': "✅ Alle Sets wurden generiert!",
        'download': "ZIP-Datei herunterladen",
        'error_api': "❌ Bitte geben Sie einen gültigen OpenAI API Key ein.",
        'sidebar_lang': "Sprache / Language",
        'lang_button': "Switch to English"
    },
    'en': {
        'title': "🧠 AI Math Teacher (Sachsen-Anhalt Expert Mode)",
        'subtitle': "Uses Artificial Intelligence to dynamically reference the Class 5 Gymnasium syllabus (Sachsen-Anhalt) and generate tough problems.",
        'api_label': "Enter OpenAI API Key:",
        'api_help': "Required to query the AI models. Not saved permanently.",
        'sets_label': "Number of problem sets:",
        'start_button': "🚀 Generate Problems with AI",
        'status_gen': "🤖 AI is analyzing the syllabus and generating Set {}/{}...",
        'success': "✅ All sets generated!",
        'download': "Download ZIP file",
        'error_api': "❌ Please enter a valid OpenAI API Key.",
        'sidebar_lang': "Language",
        'lang_button': "Zu Deutsch wechseln"
    }
}

# --- AI Generation Function ---

def generate_problems_with_ai(client, set_num, num_problems=50):
    """
    Uses OpenAI to generate a structured list of math problems based on the syllabus.
    """
    
    # The Prompt is the "Search Engine" here. It forces the AI to look up its internal 
    # knowledge of the Sachsen-Anhalt Syllabus.
    system_prompt = (
        "Du bist ein strenger Mathematik-Lehrer an einem Gymnasium in Sachsen-Anhalt. "
        "Du kennst den aktuellen Rahmenlehrplan für die Klasse 5 (Gymnasium) auswendig. "
        "Deine Aufgabe ist es, extrem anspruchsvolle Prüfungsaufgaben zu erstellen."
    )
    
    user_prompt = f"""
    Erstelle bitte {num_problems} verschiedene Mathematikaufgaben für die Klasse 5 Gymnasium (Sachsen-Anhalt).
    
    Beziehe dich STRIKT auf folgende Themenbereiche des Lehrplans:
    1. Natürliche Zahlen (Große Zahlen, Runden, Potenzen).
    2. Grundrechenarten (Schriftlich, vorteilhaftes Rechnen, Klammerregeln).
    3. Geometrie (Grundbegriffe, Koordinatensystem, Umfang/Fläche von Rechteck/Quadrat, Symmetrie).
    4. Größen (Länge, Masse, Zeit, Geld).
    5. Stochastik/Daten (einfache Diagramme, Häufigkeiten - rein textbasiert).

    Anforderungen:
    - Die Aufgaben sollen 'schwer' sein (Transferaufgaben).
    - Mische Sachaufgaben (Textaufgaben) und Rechenaufgaben.
    - Formatiere die Ausgabe als nummerierte Liste (1., 2., 3., ...).
    - Keine Markdown-Formatierung wie Fettdruck (**), nur reiner Text.
    - Gib NUR die Aufgaben zurück, keinen Einleitungstext.
    """

    try:
        response = client.chat.completions.create(
            model="gpt-4o", # Or "gpt-3.5-turbo" for lower cost
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7, # Higher creativity for "dynamic" problems
        )
        content = response.choices[0].message.content
        # Split content into lines/problems to handle them as list items
        problems = [line.strip() for line in content.split('\n') if line.strip() and (line[0].isdigit() or line.startswith('-'))]
        return problems
    except Exception as e:
        return [f"Fehler bei der AI-Generierung: {str(e)}"]

# --- Document Helper ---

def create_word_document(problems, set_number):
    document = Document()
    document.add_heading(f'Gymnasium Kl. 5 (Sachsen-Anhalt) - AI Set {set_number}', 0)
    document.add_paragraph("Generiert basierend auf dem aktuellen Rahmenlehrplan.")
    
    for problem in problems:
        # Clean up numbering if the AI included it, to let Word handle formatting or just text
        document.add_paragraph(problem)
        
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- Main App Logic ---

def set_language():
    st.session_state['lang'] = 'en' if st.session_state.get('lang', 'de') == 'de' else 'de'

def main():
    st.set_page_config(page_title="AI Math Generator", layout="centered")

    if 'lang' not in st.session_state:
        st.session_state['lang'] = 'de'
    
    t = TEXTS[st.session_state['lang']]

    # Sidebar
    st.sidebar.header(t['sidebar_lang'])
    st.sidebar.button(t['lang_button'], on_click=set_language)
    st.sidebar.markdown("---")
    
    # API Input
    api_key = st.sidebar.text_input(t['api_label'], type="password", help=t['api_help'])

    # Title
    st.title(t['title'])
    st.markdown(t['subtitle'])
    st.info("ℹ️ " + ("Dynamic Syllabus Search enabled via OpenAI GPT-4o." if st.session_state['lang']=='en' else "Dynamische Lehrplan-Suche aktiviert via OpenAI GPT-4o."))

    # Inputs
    num_sets = st.number_input(t['sets_label'], min_value=1, max_value=5, value=1)
    
    st.markdown("---")

    if st.button(t['start_button'], type="primary"):
        if not api_key:
            st.error(t['error_api'])
            return

        client = OpenAI(api_key=api_key)
        zip_buffer = BytesIO()
        
        # Status container
        with st.status("AI Working...", expanded=True) as status:
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                
                for i in range(1, int(num_sets) + 1):
                    status.write(t['status_gen'].format(i, num_sets))
                    
                    # Call AI
                    problems = generate_problems_with_ai(client, i)
                    
                    # Create Doc
                    docx_bytes = create_word_document(problems, i)
                    zf.writestr(f"AI_Math_Set_{i}.docx", docx_bytes)
                    
            zip_buffer.seek(0)
            status.update(label=t['success'], state="complete", expanded=False)
            
        st.download_button(
            label=t['download'],
            data=zip_buffer,
            file_name="AI_Sachsen_Anhalt_Math.zip",
            mime="application/zip",
            type="primary"
        )

if __name__ == "__main__":
    main()