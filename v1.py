import streamlit as st
import random
from docx import Document
from io import BytesIO
import zipfile
import os

# --- Generator Functions (Enhanced to maximize internal randomness) ---

# Global list of all generator functions
GENERATORS = []

def generate_arithmetic():
    """Generiert eine einfache Rechenaufgabe (Addition, Subtraktion, Multiplikation, Division)."""
    operator = random.choice(['+', '-', 'x', ':'])
    
    if operator in ['+', '-']:
        # Addition/Subtraktion von größeren Zahlen
        num1 = random.randint(1000, 99999)
        num2 = random.randint(100, 9999)
        if operator == '-' and num1 < num2:
            num1, num2 = num2, num1
        problem = f"{num1} {operator} {num2}"
        return f"Berechne: {problem}"
    elif operator == 'x':
        # Multiplikation
        num1 = random.randint(10, 500)
        num2 = random.randint(2, 50)
        return f"Berechne: {num1} x {num2} (schriftlich)."
    else: # Division
        # Sicherstellen, dass die Division ganzzahlig ist
        result = random.randint(10, 500)
        num2 = random.randint(2, 25)
        num1 = num2 * result
        return f"Berechne: {num1} : {num2} (schriftlich)."

GENERATORS.append(generate_arithmetic)

def generate_rounding():
    """Generiert eine Aufgabe zum Runden von Zahlen."""
    num = random.randint(100000, 9999999)
    place = random.choice(['Zehner', 'Tausender', 'Zehntausender', 'Hunderttausender'])
    return f"Runde die Zahl {num} auf die nächsten {place}."

GENERATORS.append(generate_rounding)

def generate_order_of_operations():
    """Generiert eine Punkt-vor-Strich-Aufgabe mit Klammern."""
    a = random.randint(2, 12)
    b = random.randint(2, 8)
    c = random.randint(2, 5)
    d = random.randint(5, 20)
    
    # Wähle ein zufälliges Muster
    pattern = random.choice([
        f"Löse: ({a} + {b}) x {c} - {d}",
        f"Löse: {a} x {b} + {c} : 1",
        f"Löse: {a} + {b} x ({c} + {d})",
        f"Löse: {a} x ({b} - {c}) : 2"
    ])
    return pattern

GENERATORS.append(generate_order_of_operations)

def generate_units_conversion():
    """Generiert eine Aufgabe zur Einheitenumrechnung."""
    unit_choice = random.choice(['Länge', 'Zeit', 'Masse'])
    
    if unit_choice == 'Länge':
        # cm zu m, oder km zu m
        if random.choice([True, False]):
            value = random.randint(100, 5000)
            return f"Wandle um: {value} Zentimeter (cm) in Meter (m)."
        else:
            value = round(random.uniform(1.0, 10.0), 2)
            return f"Wandle um: {value} Kilometer (km) in Meter (m)."
            
    elif unit_choice == 'Zeit':
        # h zu min, oder min zu h/min
        if random.choice([True, False]):
            h = random.randint(1, 4)
            m = random.randint(1, 59)
            return f"Wandle um: {h} Stunden (h) und {m} Minuten (min) in Gesamtminuten."
        else:
            m = random.randint(70, 300)
            return f"Wandle um: {m} Minuten (min) in Stunden (h) und Minuten (min)."
            
    else: # Masse
        # g zu kg, oder t zu kg
        if random.choice([True, False]):
            value = random.randint(500, 9000)
            return f"Wandle um: {value} Gramm (g) in Kilogramm (kg)."
        else:
            value = round(random.uniform(0.1, 3.0), 2)
            return f"Wandle um: {value} Tonnen (t) in Kilogramm (kg)."

GENERATORS.append(generate_units_conversion)

def generate_geometry_perimeter_area():
    """Generiert eine Aufgabe zu Umfang oder Fläche eines Rechtecks/Quadrats."""
    shape = random.choice(['Rechteck', 'Quadrat'])
    
    if shape == 'Quadrat':
        side = random.randint(5, 25)
        choice = random.choice(['Umfang', 'Flächeninhalt'])
        return f"Ein Quadrat hat eine Seitenlänge von {side} cm. Berechne den {choice}."
    else: # Rechteck
        length = random.randint(10, 40)
        width = random.randint(5, 20)
        choice = random.choice(['Umfang', 'Flächeninhalt'])
        return f"Ein Rechteck ist {length} m lang und {width} m breit. Wie groß ist der {choice}?"

GENERATORS.append(generate_geometry_perimeter_area)

def generate_symmetry():
    """Generiert eine Frage zur Achsensymmetrie oder Koordinaten."""
    
    if random.choice([True, False]):
        figure = random.choice(['Quadrat', 'gleichseitiges Dreieck', 'Kreis', 'Buchstabe H', 'Rechteck'])
        return f"Wie viele Symmetrieachsen besitzt ein {figure}?"
    else:
        # Koordinatensystem-Aufgabe
        x1, y1 = random.randint(1, 10), random.randint(1, 10)
        dx, dy = random.choice([-2, -1, 1, 2]), random.choice([-2, -1, 1, 2])
        return f"Ein Punkt A liegt bei ({x1}|{y1}). Er wird um {dx} Einheiten nach rechts und {dy} Einheiten nach oben verschoben. Was sind die neuen Koordinaten?"

GENERATORS.append(generate_symmetry)

def generate_word_problem():
    """Generiert eine einfache Sachaufgabe."""
    
    problem_type = random.choice(['Kosten', 'Zeitplan', 'Gesamtmenge'])
    
    if problem_type == 'Kosten':
        items = random.randint(3, 10)
        price_per_item = random.randint(1, 5)
        start_money = random.choice([20, 50, 100])
        return f"Ein Schüler kauft {items} Hefte zu je {price_per_item}€. Er bezahlt mit einem {start_money}€-Schein. Wie viel Wechselgeld erhält er?"
    elif problem_type == 'Zeitplan':
        start_h = random.randint(7, 10)
        start_m = random.choice([0, 15, 30, 45])
        duration_m = random.randint(60, 180)
        return f"Eine Zugfahrt dauert {duration_m} Minuten. Der Zug fährt um {start_h}:{start_m:02d} Uhr ab. Wann kommt er an?"
    else: # Gesamtmenge
        daily = random.randint(50, 200)
        days = random.choice([7, 30, 56])
        return f"Ein Bäcker backt täglich {daily} Brötchen. Wie viele Brötchen backt er in {days} Tagen?"

GENERATORS.append(generate_word_problem)


# --- Dokumenten- und Streamlit-Funktionen ---

def create_single_problem_set(num_problems=50):
    """Generiert eine Liste von 50 Aufgaben, indem zufällig aus allen Generatoren gezogen wird."""
    problems = []
    
    # Stellen Sie sicher, dass jede Kategorie mindestens 5 Mal vorkommt, der Rest ist zufällig
    min_per_category = 5
    required_problems = []
    
    # Füge mindestens 5 Probleme von jedem Typ hinzu
    for generator in GENERATORS:
        for _ in range(min_per_category):
            required_problems.append(generator)
            
    # Füge zusätzliche zufällige Generatoren hinzu, bis 50 erreicht sind
    num_random_fill = num_problems - len(required_problems)
    for _ in range(num_random_fill):
        required_problems.append(random.choice(GENERATORS))
        
    # Mische die Liste der Generatoren, um die Reihenfolge völlig zufällig zu machen
    random.shuffle(required_problems)
    
    # Generiere die Probleme
    for i, generator in enumerate(required_problems, 1):
        problem = generator()
        problems.append(f"{i}. {problem}")
        
    return problems


def create_word_document(problems, set_number):
    """Erstellt ein Word-Dokument (im Speicher) und gibt es als Bytes zurück."""
    document = Document()
    document.add_heading(f'Mathematikaufgaben Gymnasium Kl. 5 (Sachsen-Anhalt) - Set {set_number}', 0)
    for problem in problems:
        document.add_paragraph(problem)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

def main():
    st.set_page_config(page_title="Matheaufgaben Generator (Gymnasium 5)", layout="centered")
    st.title("🔢 Matheaufgaben Generator (Gymnasium Kl. 5)")
    st.markdown("Generiert **50 völlig zufällige und einzigartige** Übungsblätter nach dem Lehrplan Sachsen-Anhalt.")
    st.markdown("---")

    # User Inputs
    num_sets = st.number_input(
        "1. Anzahl der benötigten Aufgabensätze:", 
        min_value=1, 
        max_value=30, # Erhöht auf 30 Sets
        value=5,
        help="Wie viele verschiedene Sätze von 50 Aufgaben benötigen Sie? (Maximale Varianz)"
    )
    
    download_location = st.text_input(
        "2. Gewünschter lokaler Speicherpfad (zur Information):",
        value="C:/Users/IhrName/Downloads/Matheaufgaben/",
        help="Der Browser kann Dateien nicht direkt dorthin speichern. Der Pfad dient nur als Hinweis."
    )
    
    st.markdown("---")

    if st.button(f"Starte Generierung von {num_sets} Sätzen und erstelle ZIP-Datei"):
        st.subheader("⬇️ Generierung gestartet...")
        
        # 1. ZIP-Archiv im Speicher vorbereiten
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
            
            for i in range(1, int(num_sets) + 1):
                # Generiere und erstelle Word-Datei (als Bytes)
                with st.spinner(f"Generiere Aufgabensatz {i} (50 Aufgaben)..."):
                    # HOHE RANDOMISIERUNG: Jeder Satz ist ein Unikat, da die Generatoren zufällig gemischt werden
                    problems = create_single_problem_set(num_problems=50)
                    docx_bytes = create_word_document(problems, i)
                    filename = f"Matheaufgaben_Set_{i}.docx"
                    
                    # Füge die Word-Datei zur ZIP-Datei hinzu
                    zf.writestr(filename, docx_bytes)
                    st.success(f"✅ Aufgabensatz {i} fertiggestellt und dem ZIP hinzugefügt.")
        
        # Gehe an den Anfang des Puffers
        zip_buffer.seek(0)
        
        st.markdown("---")
        st.markdown(f"**ℹ️ Hinweis:** Die Dateien wären lokal unter dem Pfad: `{download_location}` gespeichert worden.")
        
        # 2. Streamlit Download Button für das ZIP-Archiv
        st.download_button(
            label="Alle Sätze als ZIP-Datei herunterladen",
            data=zip_buffer,
            file_name="Matheaufgaben_Klasse_5_Sets.zip",
            mime="application/zip",
            key='download_zip_button'
        )
        st.balloons()

if __name__ == "__main__":
    main()