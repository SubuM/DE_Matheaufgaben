import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO
import zipfile
import time

# --- Configuration & Texts ---
TEXTS = {
    'de': {
        'title': "🧠 Gemini Mathe-Lehrer (Sachsen-Anhalt Modus)",
        'subtitle': "Nutzt Google Gemini, um den Lehrplan der 5. Klasse Gymnasium (Sachsen-Anhalt) dynamisch zu analysieren und schwere Aufgaben zu erstellen.",
        'api_label': "Google API Key eingeben:",
        'api_help': "Der Key wird benötigt, um Gemini abzufragen. Erstellen Sie einen Key in Google AI Studio.",
        'api_loaded': "🔑 API Key sicher aus secrets.toml geladen.",
        'sets_label': "Anzahl der Aufgabensätze:",
        'start_button': "🚀 Generiere Aufgaben mit Gemini",
        'status_gen': "🤖 Gemini analysiert den Lehrplan und generiert Set {}/{}...",
        'success': "✅ Alle Sets wurden generiert!",
        'download': "ZIP-Datei herunterladen",
        'error_api': "❌ Kein API Key gefunden. Bitte in secrets.toml hinterlegen oder eingeben.",
        'sidebar_lang': "Sprache / Language",
        'lang_button': "Switch to English"
    },
    'en': {
        'title': "🧠 Gemini Math Teacher (Sachsen-Anhalt Mode)",
        'subtitle': "Uses Google Gemini to dynamically reference the Class 5 Gymnasium syllabus (Sachsen-Anhalt) and generate tough problems.",
        'api_label': "Enter Google API Key:",
        'api_help': "Required to query Gemini. Create a key in Google AI Studio.",
        'api_loaded': "🔑 API Key loaded securely from secrets.toml.",
        'sets_label': "Number of problem sets:",
        'start_button': "🚀 Generate Problems with Gemini",
        'status_gen': "🤖 Gemini is analyzing the syllabus and generating Set {}/{}...",
        'success': "✅ All sets generated!",
        'download': "Download ZIP file",
        'error_api': "❌ No API Key found. Please add to secrets.toml or enter manually.",
        'sidebar_lang': "Language",
        'lang_button': "Zu Deutsch wechseln"
    }
}

# --- Gemini Generation Function ---

def generate_problems_with_gemini(api_key, set_num, num_problems=50):
    """
    Uses Google Gemini to generate a structured list of math problems.
    """
    genai.configure(api_key=api_key)
    
    # Use 'gemini-1.5-flash' for speed or 'gemini-1.5-pro' for complex reasoning
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    Du bist ein strenger Mathematik-Lehrer an einem Gymnasium in Sachsen-Anhalt.
    Erstelle bitte ein Übungsblatt mit exakt {num_problems} verschiedenen Mathematikaufgaben für die Klasse 5.
    
    Nutze dein Wissen über den offiziellen Rahmenlehrplan für Gymnasium in Sachsen-Anhalt.
    
    Beziehe dich STRIKT auf folgende Themenbereiche:
    1. Natürliche Zahlen (Große Zahlen, Runden, Potenzen).
    2. Grundrechenarten (Schriftlich, vorteilhaftes Rechnen, Klammerregeln - Punkt vor Strich).
    3. Geometrie (Grundbegriffe, Koordinatensystem, Umfang/Fläche von Rechteck/Quadrat, Symmetrie).
    4. Größen (Länge, Masse, Zeit, Geld).
    5. Stochastik/Daten (einfache Diagramme interpretieren).

    Anforderungen:
    - Die Aufgaben sollen 'schwer' sein (Transferaufgaben, Sachaufgaben).
    - Formatiere die Ausgabe als einfache nummerierte Liste (1. Aufgabe...).
    - Keine Markdown-Formatierung wie Fettdruck (**), nur reiner Text.
    - Gib NUR die Aufgaben zurück, keine Einleitung oder Schlussworte.
    """

    try:
        response = model.generate_content(prompt)
        content = response.text
        problems = [line.strip() for line in content.split('\n') if line.strip() and (line[0].isdigit() or line.startswith('-'))]
        if not problems:
            problems = content.split('\n')
        return problems
    except Exception as e:
        return [f"Fehler bei der Gemini-Generierung: {str(e)}"]

# --- Document Helper ---

def create_word_document(problems, set_number):
    document = Document()
    document.add_heading(f'Set {set_number}', 0)
    document.add_paragraph("Generiert mit Google Gemini basierend auf dem Rahmenlehrplan.")
    for problem in problems:
        document.add_paragraph(problem)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- Main App Logic ---

def set_language():
    st.session_state['lang'] = 'en' if st.session_state.get('lang', 'de') == 'de' else 'de'

def main():
    st.set_page_config(page_title="Gemini Math Generator", layout="centered")

    if 'lang' not in st.session_state:
        st.session_state['lang'] = 'de'
    
    t = TEXTS[st.session_state['lang']]

    # Sidebar
    st.sidebar.header(t['sidebar_lang'])
    st.sidebar.button(t['lang_button'], on_click=set_language)
    st.sidebar.markdown("---")
    
    # --- API Key Logic: Check Secrets First ---
    api_key = None
    
    # Check if 'genai' section exists in secrets and has 'api_key'
    if "genai" in st.secrets and "api_key" in st.secrets["genai"]:
        api_key = st.secrets["genai"]["api_key"]
        st.sidebar.success(t['api_loaded'])
    else:
        # Fallback to manual input if secret is missing
        api_key = st.sidebar.text_input(t['api_label'], type="password", help=t['api_help'])

    # Title
    st.title(t['title'])
    st.markdown(t['subtitle'])
    st.info("ℹ️ " + ("Powered by Google Gemini 1.5 Flash" if st.session_state['lang']=='en' else "Angetrieben von Google Gemini 1.5 Flash"))

    # Inputs
    num_sets = st.number_input(t['sets_label'], min_value=1, max_value=5, value=1)
    
    st.markdown("---")

    if st.button(t['start_button'], type="primary"):
        if not api_key:
            st.error(t['error_api'])
            return

        zip_buffer = BytesIO()
        
        # Status container
        with st.status("Gemini Working...", expanded=True) as status:
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                
                for i in range(1, int(num_sets) + 1):
                    status.write(t['status_gen'].format(i, num_sets))
                    
                    # Call Gemini
                    problems = generate_problems_with_gemini(api_key, i)
                    
                    # Create Doc
                    docx_bytes = create_word_document(problems, i)
                    zf.writestr(f"Gemini_Math_Set_{i}.docx", docx_bytes)
                    
            zip_buffer.seek(0)
            status.update(label=t['success'], state="complete", expanded=False)
            
        st.download_button(
            label=t['download'],
            data=zip_buffer,
            file_name="Gemini_Sachsen_Anhalt_Math.zip",
            mime="application/zip",
            type="primary"
        )

if __name__ == "__main__":
    main()