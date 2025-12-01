import streamlit as st
import random
from docx import Document
from io import BytesIO
import zipfile
import os
import time # Used for simulated processing time/clearer status updates

# --- Generator Functions (Tough Problems - DEFINED HERE) ---

GENERATORS = []

# --- Tough Arithmetic ---
def generate_arithmetic_tough():
    operator = random.choice(['+', '-', 'x', ':'])
    if operator in ['+', '-']:
        num1 = random.randint(100000, 999999999)
        num2 = random.randint(10000, 9999999)
        num3 = random.randint(100, 5000)
        pattern = random.choice([
            f"Berechne: {num1} {operator} {num2} + {num3}",
            f"Berechne: {num1} + {num2} - {num3}"
        ])
        return pattern
    elif operator == 'x':
        num1 = random.randint(100, 999)
        num2 = random.randint(10, 999)
        return f"Berechne schriftlich: {num1} x {num2}."
    else:
        divisor = random.randint(11, 25)
        quotient = random.randint(500, 2000)
        num1 = divisor * quotient + random.randint(0, divisor - 1)
        return f"Berechne schriftlich: {num1} : {divisor} (mit Rest)."
GENERATORS.append(generate_arithmetic_tough)

# --- Tough Rounding ---
def generate_rounding_tough():
    num = random.randint(10000000, 999999999)
    place = random.choice(['Millionen', 'Zehnmillionen', 'Hunderttausender'])
    return f"Runde die Zahl {num} auf die nächsten {place}."
GENERATORS.append(generate_rounding_tough)

# --- Tough Order of Operations ---
def generate_order_of_operations_tough():
    """Zweifache Klammerung und Potenzierung."""
    a = random.randint(2, 5)
    b = random.randint(5, 15)
    c = random.randint(2, 5)
    d = random.randint(10, 30)
    e = random.randint(1, 3) # Exponent
    
    pattern = random.choice([
        f"Löse: ({b} - {c} + {d}) x {a}",
        f"Löse: {d} + [{b} x ({c} + {a})]",
        # FIX: Removed LaTeX and used standard Python exponent notation (**) for clarity
        f"Löse: {a}**{e} + {b} x ({d} - {c})" 
    ])
    return pattern
GENERATORS.append(generate_order_of_operations_tough)

# --- Tough Units Conversion ---
def generate_units_conversion_tough():
    unit_choice = random.choice(['Länge', 'Zeit', 'Masse'])
    if unit_choice == 'Länge':
        km = random.randint(5, 50)
        m = random.randint(1, 999)
        cm = random.randint(1, 99)
        return f"Wandle um: {km} km, {m} m und {cm} cm in Gesamtmetern (m)."
    elif unit_choice == 'Zeit':
        h = random.randint(4, 10)
        m = random.randint(1, 59)
        s = random.randint(1, 59)
        return f"Wandle um: {h} h, {m} min und {s} s in Gesamtsekunden (s)."
    else:
        t = random.randint(1, 5)
        kg = random.randint(10, 999)
        g = random.randint(1, 999)
        return f"Wandle um: {t} t, {kg} kg und {g} g in Gesamtgramm (g)."
GENERATORS.append(generate_units_conversion_tough)

# --- Tough Geometry/Area ---
def generate_geometry_perimeter_area_tough():
    """Rückwärtsaufgaben oder Aufgaben mit zusammengesetzten Figuren."""
    choice = random.choice(['Fläche_Rück', 'Umfang_Rück', 'Zusammengesetzt'])
    
    if choice == 'Fläche_Rück':
        area = random.randint(100, 500)
        width = random.randint(5, 20)
        # FIX: Replaced LaTeX $\text{cm}^2$ with Unicode cm²
        return f"Die Fläche eines Rechtecks beträgt {area} cm². Die Breite ist {width} cm. Berechne die Länge und den Umfang."
    elif choice == 'Umfang_Rück':
        perimeter = random.randint(80, 200)
        length = random.randint(20, 50)
        return f"Der Umfang eines Rechtecks ist {perimeter} m. Die Länge ist {length} m. Berechne die Breite und die Fläche."
    else: # Zusammengesetzt
        l1 = random.randint(10, 20)
        w1 = random.randint(5, 10)
        l2 = random.randint(5, 10)
        w2 = random.randint(2, 5)
        return f"Eine L-förmige Figur besteht aus zwei Rechtecken: R1 ({l1}x{w1} cm) und R2 ({l2}x{w2} cm). Berechne den Gesamtflächeninhalt."
GENERATORS.append(generate_geometry_perimeter_area_tough)

# --- Tough Symmetry/Coordinates ---
def generate_symmetry_tough():
    if random.choice([True, False]):
        figure = random.choice(['gleichschenkliges Trapez', 'Rhombus'])
        return f"Zeichnen Sie ein {figure} und bestimmen Sie die Anzahl seiner Symmetrieachsen."
    else:
        x, y = random.randint(1, 5), random.randint(1, 5)
        return f"Der Punkt P({x}|{y}) wird an der y-Achse gespiegelt. Geben Sie die neuen Koordinaten P' an."
GENERATORS.append(generate_symmetry_tough)

# --- Tough Word Problem ---
def generate_word_problem_tough():
    item_count = random.randint(5, 15)
    price = random.randint(2, 8)
    weight_g = random.randint(100, 500)
    return f"Ein Händler kauft {item_count} Kisten Äpfel zu je {price}€ pro Kiste. Jede Kiste wiegt {weight_g} g. Wie viel bezahlt er insgesamt und wie schwer sind alle Kisten zusammen in Kilogramm?"
GENERATORS.append(generate_word_problem_tough)

# --- Document and Helper Functions (Unchanged Logic) ---

def create_single_problem_set(num_problems=50):
    problems = []
    min_per_category = 7 
    required_problems = []
    
    for generator in GENERATORS:
        for _ in range(min_per_category):
            required_problems.append(generator)
            
    num_random_fill = num_problems - len(required_problems)
    for _ in range(num_random_fill):
        required_problems.append(random.choice(GENERATORS))
        
    random.shuffle(required_problems)
    
    for i, generator in enumerate(required_problems, 1):
        problem = generator()
        problems.append(f"{i}. {problem}")
        
    return problems

def create_word_document(problems, set_number):
    document = Document()
    document.add_heading(f'Schwere Mathematikaufgaben Gymnasium Kl. 5 (Sachsen-Anhalt) - Set {set_number}', 0)
    document.add_paragraph("Dies sind Übungen mit erhöhtem Schwierigkeitsgrad.")
    for problem in problems:
        document.add_paragraph(problem)
    bio = BytesIO()
    document.save(bio)
    bio.seek(0)
    return bio.getvalue()

# --- Streamlit Login and Main Program ---

def login_form():
    """Zeigt das Login-Formular an und prüft die Anmeldedaten aus st.secrets."""
    
    try:
        STATIC_USER = st.secrets["auth"]["username"]
        STATIC_PASS = st.secrets["auth"]["password"]
    except KeyError:
        st.error("❌ Konfigurationsfehler: Bitte stellen Sie sicher, dass die Datei '.streamlit/secrets.toml' mit der Sektion [auth] existiert.")
        return

    st.sidebar.header("Login erforderlich")
    with st.sidebar.form("login_form"):
        username = st.text_input("Benutzername")
        password = st.text_input("Passwort", type="password")
        submitted = st.form_submit_button("Login")
        
        if submitted:
            if username == STATIC_USER and password == STATIC_PASS:
                st.session_state["authenticated"] = True
                st.success("Login erfolgreich!")
            else:
                st.error("Ungültiger Benutzername oder Passwort.")

def main_app():
    """Die Hauptanwendung, die nur nach erfolgreichem Login angezeigt wird."""
    st.title("💪 Matheaufgaben Generator (Hoher Schwierigkeitsgrad)")
    st.markdown("Generiert bis zu **10 völlig zufällige und schwere** Übungsblätter.")
    st.markdown("---")

    # User Inputs
    num_sets = st.number_input(
        "1. Anzahl der benötigten Aufgabensätze:", 
        min_value=1, 
        max_value=10, 
        value=5,
        help="Maximal 10 Sets empfohlen, um die Einzigartigkeit zu garantieren."
    )
    
    download_location = st.text_input(
        "2. Gewünschter lokaler Speicherpfad (zur Information):",
        value="C:/Users/IhrName/Downloads/Matheaufgaben_Schwer/",
        help="Der Pfad dient nur als Hinweis, der Download erfolgt über den Browser."
    )
    
    st.markdown("---")

    if st.button(f"Starte Generierung von {num_sets} Sätzen und erstelle ZIP-Datei"):
        
        # 1. ZIP-Archiv im Speicher vorbereiten
        zip_buffer = BytesIO()
        
        # NEU: Verwende st.status für klare visuelle Rückmeldung
        with st.status("🛠️ Starte die Generierung...", expanded=True) as status:
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                
                for i in range(1, int(num_sets) + 1):
                    # Statusmeldung für jedes Set
                    status.write(f"➡️ **Set {i}/{num_sets}:** Erstelle 50 schwere Aufgaben...")
                    
                    # Generiere und erstelle Word-Datei (als Bytes)
                    problems = create_single_problem_set(num_problems=50)
                    docx_bytes = create_word_document(problems, i)
                    filename = f"Matheaufgaben_Set_SCHWER_{i}.docx"
                    
                    # Füge die Word-Datei zur ZIP-Datei hinzu
                    zf.writestr(filename, docx_bytes)
                    status.write(f"✅ **Set {i}/{num_sets}:** Erfolgreich zur ZIP-Datei hinzugefügt.")
                
                status.update(label="📦 Komprimiere und finalisiere das ZIP-Archiv...", state="running")
                time.sleep(1) # Simulate compression time
            
            # Schließe den ZIP-Vorgang im Buffer ab
            zip_buffer.seek(0)
            
            # Finale Statusmeldung
            status.update(label=f"🎉 {num_sets} Aufgabensätze sind bereit zum Download!", state="complete", expanded=False)
        
        st.markdown("---")
        st.markdown(f"**ℹ️ Hinweis:** Die Dateien wären lokal unter dem Pfad: `{download_location}` gespeichert worden.")
        
        # 2. Streamlit Download Button für das ZIP-Archiv
        st.download_button(
            label="Alle Sätze als ZIP-Datei herunterladen",
            data=zip_buffer,
            file_name="Matheaufgaben_Klasse_5_SCHWER.zip",
            mime="application/zip",
            key='download_zip_button'
        )

# Hauptfunktion, die den Login-Zustand verwaltet
if __name__ == "__main__":
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False

    if st.session_state["authenticated"]:
        main_app()
    else:
        login_form()